"""Generates the fictional Contoso Corp policy corpus for Phase 1.

Single source of truth: the DOCUMENTS list below holds each document's title,
format, and nested section/subsection headings + prose. render_markdown /
render_docx / render_pdf turn that structure into real files with real
heading hierarchy (Word heading styles for .docx, font-size-differentiated
styles for .pdf, '#'/'##'/'###' for .md) -- not just visually bold text --
so a later structure-aware chunker has real section boundaries to key off
of. manifest.json is derived from the same DOCUMENTS list, so the outline
recorded there can never drift from what's actually in the files.

Regenerate with:
    cd corpus && python3 -m venv .venv && ./.venv/bin/pip install -r requirements.txt
    ./.venv/bin/python generate_corpus.py
"""

import json
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUT_DIR = Path(__file__).parent

# ---------------------------------------------------------------------------
# Document content. Each "section" may have its own "paragraphs" and, one
# level deeper, "subsections" (which have their own "paragraphs"). That is
# the full depth used across the corpus: Title > Section > Subsection.
# ---------------------------------------------------------------------------

DOCUMENTS = [
    {
        "key": "employee-handbook",
        "title": "Contoso Corp Employee Handbook",
        "doc_type": "pdf",
        "filename": "employee-handbook.pdf",
        "sections": [
            {
                "heading": "Welcome to Contoso Corp",
                "paragraphs": [
                    "Contoso Corp is a technology company headquartered in Northlake, Washington, with "
                    "approximately 4,800 employees across engineering, sales, operations, and support "
                    "functions. This handbook explains the policies, benefits, and expectations that "
                    "apply to every Contoso employee, and points you to the standalone policy documents "
                    "that cover each topic in more depth.",
                    "Our mission is to help our customers run their operations more reliably, and we "
                    "expect every employee to treat that mission, our customers, and each other with "
                    "integrity, respect, and accountability. Specific expectations for workplace conduct "
                    "are detailed in the Contoso Corp Code of Conduct and Business Ethics Policy.",
                ],
            },
            {
                "heading": "Employment Classifications",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Full-Time and Part-Time Status",
                        "paragraphs": [
                            "Employees who are regularly scheduled for 30 or more hours per week are "
                            "classified as full-time and are eligible for the full benefits package "
                            "described in the Contoso Corp Health and Welfare Benefits Guide. Employees "
                            "scheduled for fewer than 30 hours per week are classified as part-time and "
                            "are eligible for a prorated subset of benefits.",
                        ],
                    },
                    {
                        "heading": "Exempt and Non-Exempt Status",
                        "paragraphs": [
                            "Exempt employees are paid a fixed salary and are not eligible for overtime "
                            "pay. Non-exempt employees are paid hourly and are eligible for overtime pay "
                            "in accordance with the Contoso Corp Compensation and Pay Practices Policy. "
                            "Your offer letter states which classification applies to your role.",
                        ],
                    },
                    {
                        "heading": "Introductory Period",
                        "paragraphs": [
                            "New employees complete a 90-day introductory period during which their "
                            "manager provides regular feedback on role expectations. Completion of the "
                            "introductory period does not change employment status; Contoso employment "
                            "is at-will unless a signed agreement states otherwise.",
                        ],
                    },
                ],
            },
            {
                "heading": "Workplace Conduct Expectations",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Attendance and Punctuality",
                        "paragraphs": [
                            "Employees are expected to be reliably available during their scheduled "
                            "hours, whether working on-site or remotely. If you are unable to work as "
                            "scheduled, notify your manager as far in advance as possible. Repeated "
                            "unexplained absences are addressed through the performance management "
                            "process described in the Contoso Corp Performance Management and Review "
                            "Policy.",
                        ],
                    },
                    {
                        "heading": "Dress Code",
                        "paragraphs": [
                            "Contoso maintains a business-casual dress code for employees working from "
                            "an office location, with adjustments for roles that require safety or "
                            "client-facing attire. Remote employees have no fixed dress code but should "
                            "dress appropriately for video meetings.",
                        ],
                    },
                ],
            },
            {
                "heading": "Where to Find More Information",
                "paragraphs": [
                    "This handbook is an overview. Detailed rules for each topic live in their own "
                    "policy documents, including the Paid Time Off and Leave Policy, the Remote Work "
                    "and Flexible Work Arrangements Policy, the Information Security and Acceptable Use "
                    "Policy, and the Workplace Health and Safety Policy. If a specific question is not "
                    "answered in this handbook, check the relevant standalone policy before assuming an "
                    "answer, or contact HR at hr@contoso-corp.example.",
                ],
            },
        ],
    },
    {
        "key": "code-of-conduct",
        "title": "Contoso Corp Code of Conduct and Business Ethics Policy",
        "doc_type": "pdf",
        "filename": "code-of-conduct.pdf",
        "sections": [
            {
                "heading": "Purpose and Scope",
                "paragraphs": [
                    "This policy sets the standard of ethical conduct expected of every Contoso Corp "
                    "employee, contractor, and officer in all business dealings, whether with customers, "
                    "suppliers, competitors, or each other. It applies globally across every Contoso "
                    "location and subsidiary.",
                ],
            },
            {
                "heading": "Core Ethical Principles",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Integrity in Business Dealings",
                        "paragraphs": [
                            "Employees must never misrepresent Contoso's products, financial results, or "
                            "capabilities to customers, investors, or regulators. Any employee who "
                            "discovers a misstatement in a customer-facing or financial document must "
                            "report it immediately through the channels described below.",
                        ],
                    },
                    {
                        "heading": "Conflicts of Interest",
                        "paragraphs": [
                            "Employees must disclose to their manager and to HR any outside financial "
                            "interest, board seat, or family relationship that could reasonably influence, "
                            "or appear to influence, a business decision they make on Contoso's behalf. "
                            "Disclosure does not automatically bar the activity; it allows Contoso to "
                            "assess and, if needed, put safeguards in place.",
                        ],
                    },
                    {
                        "heading": "Gifts and Entertainment",
                        "paragraphs": [
                            "Employees may accept or offer business gifts and entertainment of modest "
                            "value (generally under $100) when doing so is customary and does not create "
                            "an appearance of impropriety. Cash gifts, and any gift offered with the "
                            "expectation of a specific business decision in return, are never permitted.",
                        ],
                    },
                ],
            },
            {
                "heading": "Reporting Violations",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Internal Reporting Channels",
                        "paragraphs": [
                            "Concerns can be raised with your manager, with HR, or with the Ethics "
                            "Office at ethics@contoso-corp.example. Reports may be made by name or, where "
                            "local law permits, anonymously.",
                        ],
                    },
                    {
                        "heading": "Non-Retaliation Commitment",
                        "paragraphs": [
                            "Contoso prohibits retaliation against any employee who reports a good-faith "
                            "concern or participates in an investigation. Reports of retaliation are "
                            "treated as seriously as the underlying conduct being reported.",
                        ],
                    },
                ],
            },
            {
                "heading": "Consequences of Violations",
                "paragraphs": [
                    "Confirmed violations of this policy may result in disciplinary action up to and "
                    "including termination of employment, and in some cases referral to law enforcement "
                    "or regulators. The severity of the response depends on the nature of the violation "
                    "and whether it was disclosed voluntarily.",
                ],
            },
        ],
    },
    {
        "key": "eeo-anti-harassment",
        "title": "Contoso Corp Equal Employment Opportunity and Anti-Harassment Policy",
        "doc_type": "pdf",
        "filename": "eeo-anti-harassment-policy.pdf",
        "sections": [
            {
                "heading": "Equal Employment Opportunity Statement",
                "paragraphs": [
                    "Contoso Corp provides equal employment opportunity to all employees and applicants "
                    "without regard to race, color, religion, sex, national origin, age, disability, "
                    "genetic information, veteran status, sexual orientation, gender identity, or any "
                    "other status protected by applicable law. This commitment applies to hiring, "
                    "promotion, compensation, and every other term of employment.",
                ],
            },
            {
                "heading": "Prohibited Harassment",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Sexual Harassment",
                        "paragraphs": [
                            "Unwelcome sexual advances, requests for sexual favors, and other verbal or "
                            "physical conduct of a sexual nature are prohibited when they create a hostile "
                            "work environment or are made a condition of employment or advancement. This "
                            "applies to conduct by coworkers, managers, and third parties such as vendors "
                            "or customers.",
                        ],
                    },
                    {
                        "heading": "Other Forms of Harassment",
                        "paragraphs": [
                            "Harassment based on any protected characteristic, including offensive jokes, "
                            "slurs, or exclusionary behavior, is prohibited regardless of whether it rises "
                            "to the level of a hostile work environment under the law. Contoso holds "
                            "itself to a higher standard than the legal minimum.",
                        ],
                    },
                ],
            },
            {
                "heading": "Reporting and Investigation Process",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "How to File a Complaint",
                        "paragraphs": [
                            "Complaints may be filed with HR, with any manager, or through the "
                            "confidential HR case portal. You do not need to raise a concern with the "
                            "person involved before filing a complaint.",
                        ],
                    },
                    {
                        "heading": "Investigation Timeline",
                        "paragraphs": [
                            "HR aims to complete a preliminary assessment within 5 business days of a "
                            "complaint and to conclude most investigations within 30 calendar days. "
                            "Complex investigations involving multiple witnesses may take longer, and "
                            "the complainant is updated on status throughout.",
                        ],
                    },
                ],
            },
            {
                "heading": "Non-Retaliation",
                "paragraphs": [
                    "Retaliation against anyone who reports harassment or discrimination in good faith, "
                    "or who participates in an investigation, is itself a violation of this policy and "
                    "is grounds for discipline up to termination.",
                ],
            },
        ],
    },
    {
        "key": "diversity-equity-inclusion",
        "title": "Contoso Corp Diversity, Equity, and Inclusion Policy",
        "doc_type": "docx",
        "filename": "diversity-equity-and-inclusion-policy.docx",
        "sections": [
            {
                "heading": "Our Commitment to Diversity, Equity, and Inclusion",
                "paragraphs": [
                    "Contoso Corp believes that a diverse workforce and an inclusive culture make us a "
                    "stronger company and a better partner to our customers. This policy describes the "
                    "programs and expectations that support that commitment, in addition to the legal "
                    "protections described in the Contoso Corp Equal Employment Opportunity and "
                    "Anti-Harassment Policy.",
                ],
            },
            {
                "heading": "Employee Resource Groups",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "How to Join or Start a Group",
                        "paragraphs": [
                            "Employees may join any Employee Resource Group (ERG) regardless of whether "
                            "they personally identify with the group's focus; allyship is welcomed. To "
                            "start a new ERG, an employee submits a proposal to the People team "
                            "describing the group's purpose and an initial slate of at least three "
                            "organizers.",
                        ],
                    },
                    {
                        "heading": "ERG Leadership and Funding",
                        "paragraphs": [
                            "Each recognized ERG elects rotating leadership annually and receives an "
                            "annual budget for events and programming, allocated by the People team "
                            "based on group size and planned activities.",
                        ],
                    },
                ],
            },
            {
                "heading": "Inclusive Hiring Practices",
                "paragraphs": [
                    "Hiring panels for open roles are required to include at least one interviewer from "
                    "outside the immediate hiring team, and job postings are reviewed for inclusive "
                    "language before publication. Structured interview scorecards are used for every "
                    "role to reduce bias in hiring decisions.",
                ],
            },
            {
                "heading": "Diversity Education and Training",
                "paragraphs": [
                    "All employees complete an unconscious bias training module within their first 90 "
                    "days, and managers complete an additional inclusive leadership module before taking "
                    "on their first direct report. These trainings are assigned automatically through "
                    "the Learning Portal.",
                ],
            },
        ],
    },
    {
        "key": "pto-leave-policy",
        "title": "Contoso Corp Paid Time Off and Leave Policy",
        "doc_type": "docx",
        "filename": "pto-and-leave-policy.docx",
        "sections": [
            {
                "heading": "Overview of Leave Types",
                "paragraphs": [
                    "This policy covers Paid Time Off (PTO), sick leave, bereavement leave, and jury "
                    "duty and civic leave for regular full-time and part-time employees. Parental and "
                    "family leave are covered separately in the Contoso Corp Parental and Family Leave "
                    "Policy.",
                ],
            },
            {
                "heading": "Paid Time Off (PTO)",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Accrual Rates",
                        "paragraphs": [
                            "Full-time employees accrue PTO at 1.67 days per month during their first "
                            "three years of service (20 days annually), increasing to 2.08 days per month "
                            "(25 days annually) beginning in year four. Part-time employees accrue PTO on "
                            "a prorated basis matching their scheduled hours.",
                        ],
                    },
                    {
                        "heading": "Requesting and Approving PTO",
                        "paragraphs": [
                            "PTO requests are submitted through the HR system at least two weeks in "
                            "advance where possible, and are approved or denied by your manager based on "
                            "team coverage. Requests for single-day absences submitted with less notice "
                            "are considered but are not guaranteed approval.",
                        ],
                    },
                    {
                        "heading": "Carryover and Payout",
                        "paragraphs": [
                            "Employees may carry over up to 5 unused PTO days into the following calendar "
                            "year; any balance above that is forfeited on December 31 unless local law "
                            "requires otherwise. Accrued, unused PTO is paid out upon separation from the "
                            "company, in accordance with the Contoso Corp Offboarding and Separation "
                            "Policy.",
                        ],
                    },
                ],
            },
            {
                "heading": "Sick Leave",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Accrual and Usage",
                        "paragraphs": [
                            "Employees accrue paid sick leave separately from PTO, at a rate of 1 hour "
                            "per 30 hours worked, up to a maximum balance of 72 hours. Sick leave may be "
                            "used for the employee's own illness or to care for an immediate family "
                            "member, and does not require advance notice.",
                        ],
                    },
                ],
            },
            {
                "heading": "Bereavement Leave",
                "paragraphs": [
                    "Employees receive up to 5 paid days of bereavement leave following the death of an "
                    "immediate family member, and up to 2 paid days for the death of a close relative "
                    "outside the immediate family. Additional unpaid leave may be arranged with your "
                    "manager.",
                ],
            },
            {
                "heading": "Jury Duty and Civic Leave",
                "paragraphs": [
                    "Contoso provides paid leave for jury duty and for appearing as a subpoenaed witness, "
                    "for the duration of the service obligation. Employees must provide a copy of the "
                    "jury summons or subpoena to HR as soon as it is received.",
                ],
            },
        ],
    },
    {
        "key": "parental-family-leave",
        "title": "Contoso Corp Parental and Family Leave Policy",
        "doc_type": "docx",
        "filename": "parental-and-family-leave-policy.docx",
        "sections": [
            {
                "heading": "Eligibility",
                "paragraphs": [
                    "Parental leave under this policy is available to any employee who has completed "
                    "at least 6 months of continuous service at the time of the qualifying birth, "
                    "adoption, or foster placement. Family and Medical Leave eligibility follows the "
                    "criteria described later in this policy.",
                ],
            },
            {
                "heading": "Birth Parent Leave",
                "paragraphs": [
                    "Birth parents receive 16 weeks of paid leave, which may begin up to 2 weeks before "
                    "the expected due date and may be taken continuously or, with manager approval, in "
                    "two blocks within the first 12 months following birth.",
                ],
            },
            {
                "heading": "Non-Birth Parent and Adoption Leave",
                "paragraphs": [
                    "Non-birth parents, and employees welcoming a child through adoption or foster "
                    "placement, receive 10 weeks of paid leave, to be taken within the first 12 months "
                    "following the qualifying event.",
                ],
            },
            {
                "heading": "Family and Medical Leave (FMLA-Aligned)",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Qualifying Reasons",
                        "paragraphs": [
                            "Beyond the paid parental leave above, employees may take unpaid, job-protected "
                            "leave for their own serious health condition, to care for a family member "
                            "with a serious health condition, or for a qualifying military family "
                            "reason, in line with federal Family and Medical Leave Act standards.",
                        ],
                    },
                    {
                        "heading": "Job Protection and Benefits Continuation",
                        "paragraphs": [
                            "Employees on qualifying leave are restored to their same or an equivalent "
                            "position upon return, and Contoso continues to pay its share of medical "
                            "premiums during the leave period as if the employee were actively working.",
                        ],
                    },
                ],
            },
            {
                "heading": "Returning to Work",
                "paragraphs": [
                    "Employees should notify their manager and HR of their expected return date at "
                    "least 2 weeks before returning. A gradual return-to-work schedule can be arranged "
                    "with manager and HR approval where a medical provider recommends one.",
                ],
            },
        ],
    },
    {
        "key": "health-welfare-benefits",
        "title": "Contoso Corp Health and Welfare Benefits Guide",
        "doc_type": "pdf",
        "filename": "health-and-welfare-benefits-guide.pdf",
        "sections": [
            {
                "heading": "Benefits Eligibility",
                "paragraphs": [
                    "Full-time employees are eligible for benefits on the first day of the month "
                    "following their start date. Part-time employees scheduled for at least 20 hours "
                    "per week are eligible for a prorated subset of medical, dental, and vision coverage.",
                ],
            },
            {
                "heading": "Medical Plan Options",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "PPO Plan",
                        "paragraphs": [
                            "The Contoso PPO plan offers access to a broad national provider network "
                            "with no referral requirements. Contoso pays 80% of the monthly premium for "
                            "employee-only coverage and 65% for coverage that includes dependents.",
                        ],
                    },
                    {
                        "heading": "High-Deductible Health Plan with HSA",
                        "paragraphs": [
                            "The high-deductible plan carries a lower monthly premium and pairs with a "
                            "Health Savings Account. Contoso contributes $750 annually to the HSA of any "
                            "employee enrolled in this plan, deposited in the first full pay period of "
                            "the plan year.",
                        ],
                    },
                ],
            },
            {
                "heading": "Dental and Vision Coverage",
                "paragraphs": [
                    "Dental coverage includes two preventive cleanings per year at no cost, with "
                    "coinsurance applying to major procedures. Vision coverage includes an annual eye "
                    "exam and an allowance toward glasses or contact lenses every 12 months.",
                ],
            },
            {
                "heading": "Life and Disability Insurance",
                "paragraphs": [
                    "Contoso automatically provides basic life insurance equal to one times annual "
                    "salary, and short-term and long-term disability coverage, at no cost to the "
                    "employee. Supplemental life insurance for the employee, spouse, or children can be "
                    "purchased during open enrollment.",
                ],
            },
            {
                "heading": "Employee Assistance Program",
                "paragraphs": [
                    "All employees and their household members have access to the Employee Assistance "
                    "Program, which provides confidential counseling sessions, legal consultation "
                    "referrals, and financial planning resources at no direct cost.",
                ],
            },
        ],
    },
    {
        "key": "retirement-savings-plan",
        "title": "Contoso Corp Retirement Savings Plan Guide",
        "doc_type": "pdf",
        "filename": "retirement-savings-plan-guide.pdf",
        "sections": [
            {
                "heading": "Plan Overview",
                "paragraphs": [
                    "The Contoso Retirement Savings Plan is a 401(k)-style defined contribution plan "
                    "administered by a third-party recordkeeper, offering a range of index and target-date "
                    "fund options.",
                ],
            },
            {
                "heading": "Eligibility and Enrollment",
                "paragraphs": [
                    "All employees are eligible to enroll starting on their first day of employment. "
                    "Employees who do not actively enroll are automatically defaulted into the plan at a "
                    "3% contribution rate beginning on their 60th day, unless they opt out.",
                ],
            },
            {
                "heading": "Contribution Limits",
                "paragraphs": [
                    "Employees may contribute up to the annual IRS elective deferral limit through pre-tax "
                    "or Roth payroll deductions. Contribution elections can be changed at any time through "
                    "the recordkeeper's website, effective the following pay period.",
                ],
            },
            {
                "heading": "Company Matching Contributions",
                "paragraphs": [
                    "Contoso matches 100% of the first 4% of eligible pay that an employee contributes, "
                    "deposited each pay period rather than as a single annual lump sum.",
                ],
            },
            {
                "heading": "Vesting Schedule",
                "paragraphs": [
                    "Employee contributions are always fully vested. Company matching contributions vest "
                    "over 3 years on a graded schedule: 33% after one year of service, 66% after two "
                    "years, and 100% after three years.",
                ],
            },
        ],
    },
    {
        "key": "remote-work-policy",
        "title": "Contoso Corp Remote Work and Flexible Work Arrangements Policy",
        "doc_type": "markdown",
        "filename": "remote-work-and-flexible-work-policy.md",
        "sections": [
            {
                "heading": "Eligibility for Remote Work",
                "paragraphs": [
                    "Roles that do not require regular on-site presence for equipment, lab access, or "
                    "in-person customer interaction are eligible for fully remote or hybrid arrangements, "
                    "subject to manager approval. Eligibility is determined by role, not tenure.",
                ],
            },
            {
                "heading": "Hybrid Work Expectations",
                "paragraphs": [
                    "Employees on a hybrid arrangement are expected to be on-site on the days agreed "
                    "with their manager, typically 2 to 3 days per week for team-based roles. Teams "
                    "publish their in-office days so collaboration time is predictable.",
                ],
            },
            {
                "heading": "Home Office Equipment and Stipend",
                "paragraphs": [
                    "Contoso provides a laptop, monitor, and standard peripherals for remote and hybrid "
                    "employees, and a one-time $300 home office setup stipend during the first 90 days "
                    "in a remote-eligible role. Ongoing home internet costs are the employee's "
                    "responsibility.",
                ],
            },
            {
                "heading": "Core Collaboration Hours",
                "paragraphs": [
                    "Regardless of location, employees are expected to be reachable during core hours of "
                    "10:00 a.m. to 3:00 p.m. in their team's primary time zone, to support real-time "
                    "collaboration across the organization.",
                ],
            },
        ],
    },
    {
        "key": "travel-expense-policy",
        "title": "Contoso Corp Business Travel and Expense Reimbursement Policy",
        "doc_type": "docx",
        "filename": "business-travel-and-expense-policy.docx",
        "sections": [
            {
                "heading": "Booking Business Travel",
                "paragraphs": [
                    "All flights and hotels for business travel are booked through the Contoso corporate "
                    "travel portal, which negotiates preferred rates. Travel booked outside the portal is "
                    "reimbursed only up to what the portal's comparable rate would have cost.",
                ],
            },
            {
                "heading": "Reimbursable Expenses",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Airfare and Lodging",
                        "paragraphs": [
                            "Economy-class airfare and standard hotel accommodations are reimbursable for "
                            "all employees. Upgrades to premium economy are reimbursable for flights over "
                            "6 hours with manager pre-approval.",
                        ],
                    },
                    {
                        "heading": "Meals and Incidentals",
                        "paragraphs": [
                            "Meals while traveling are reimbursed up to $75 per day domestically and $100 "
                            "per day internationally, itemized receipts required for any single expense "
                            "over $25. Alcohol is reimbursable only when part of a client meal.",
                        ],
                    },
                    {
                        "heading": "Personal Vehicle Mileage",
                        "paragraphs": [
                            "Employees who use a personal vehicle for business travel are reimbursed at "
                            "the current IRS standard mileage rate. Normal commuting mileage between home "
                            "and your regular office is never reimbursable.",
                        ],
                    },
                ],
            },
            {
                "heading": "Non-Reimbursable Expenses",
                "paragraphs": [
                    "Personal entertainment, travel companion costs, traffic or parking fines, and "
                    "airline club memberships are not reimbursable under any circumstances.",
                ],
            },
            {
                "heading": "Submitting Expense Reports",
                "paragraphs": [
                    "Expense reports must be submitted through the Contoso expense system within 30 days "
                    "of the trip, with itemized receipts attached for every expense over $25. Reports "
                    "submitted after 90 days require VP-level approval to be reimbursed.",
                ],
            },
        ],
    },
    {
        "key": "infosec-acceptable-use",
        "title": "Contoso Corp Information Security and Acceptable Use Policy",
        "doc_type": "markdown",
        "filename": "information-security-and-acceptable-use-policy.md",
        "sections": [
            {
                "heading": "Purpose",
                "paragraphs": [
                    "This policy defines the security standards that apply to every Contoso-owned "
                    "device, account, and network connection, and to any personal device used to access "
                    "Contoso systems.",
                ],
            },
            {
                "heading": "Acceptable Use of Company Systems",
                "paragraphs": [
                    "Company systems are provided primarily for business use. Limited, incidental "
                    "personal use is permitted as long as it does not interfere with work, consume "
                    "significant resources, or involve illegal or offensive content.",
                ],
            },
            {
                "heading": "Password and Authentication Requirements",
                "paragraphs": [
                    "All Contoso accounts require multi-factor authentication and a passphrase of at "
                    "least 14 characters, changed if there is any indication of compromise. Password "
                    "reuse across personal and Contoso accounts is prohibited.",
                ],
            },
            {
                "heading": "Handling Confidential Data",
                "paragraphs": [
                    "Confidential and customer data must be stored only in approved Contoso systems, "
                    "never copied to personal cloud storage or removable media, and encrypted in transit. "
                    "Detailed data classification rules are in the Contoso Corp Data Privacy and "
                    "Confidentiality Policy.",
                ],
            },
            {
                "heading": "Reporting Security Incidents",
                "paragraphs": [
                    "Any suspected phishing email, lost device, or unauthorized access must be reported "
                    "immediately to security@contoso-corp.example. Do not wait to confirm the incident is "
                    "real before reporting it; the security team will triage.",
                ],
            },
        ],
    },
    {
        "key": "data-privacy-confidentiality",
        "title": "Contoso Corp Data Privacy and Confidentiality Policy",
        "doc_type": "docx",
        "filename": "data-privacy-and-confidentiality-policy.docx",
        "sections": [
            {
                "heading": "Scope",
                "paragraphs": [
                    "This policy governs how Contoso employees handle confidential company information, "
                    "employee personal data, and customer data encountered in the course of their work.",
                ],
            },
            {
                "heading": "Categories of Protected Data",
                "paragraphs": [
                    "Protected data includes customer account information, unreleased product plans, "
                    "employee compensation and health data, and financial results before public "
                    "disclosure. Each category carries its own handling requirements described in the "
                    "internal data classification standard.",
                ],
            },
            {
                "heading": "Employee Confidentiality Obligations",
                "paragraphs": [
                    "Employees may access confidential data only as needed for their role, and must not "
                    "discuss it outside of Contoso, including with family members or on social media. "
                    "This obligation continues after employment ends.",
                ],
            },
            {
                "heading": "Customer Data Handling",
                "paragraphs": [
                    "Customer data may be used only for the purpose the customer authorized, and access "
                    "is logged and periodically audited. Any request from a customer to delete their data "
                    "is routed to the Privacy Office rather than handled ad hoc by individual employees.",
                ],
            },
            {
                "heading": "Data Retention and Disposal",
                "paragraphs": [
                    "Data is retained only as long as required for business or legal purposes, per the "
                    "records retention schedule maintained by Legal, and is disposed of using secure "
                    "deletion methods once that period ends.",
                ],
            },
        ],
    },
    {
        "key": "social-media-external-comms",
        "title": "Contoso Corp Social Media and External Communications Policy",
        "doc_type": "markdown",
        "filename": "social-media-and-external-communications-policy.md",
        "sections": [
            {
                "heading": "Purpose",
                "paragraphs": [
                    "This policy sets expectations for how employees use social media, both personally "
                    "and on behalf of Contoso, to protect the company's reputation and confidential "
                    "information while respecting employees' right to participate in public "
                    "conversation.",
                ],
            },
            {
                "heading": "Personal Social Media Use",
                "paragraphs": [
                    "Employees may identify their employer on personal social media profiles and are "
                    "free to discuss their work generally, but must not share confidential information "
                    "covered by the Contoso Corp Data Privacy and Confidentiality Policy or speak as if "
                    "authorized to represent Contoso's official position. Adding a standard disclaimer "
                    "that opinions are one's own is encouraged but not required.",
                ],
            },
            {
                "heading": "Speaking on Behalf of Contoso",
                "paragraphs": [
                    "Only employees specifically designated by Corporate Communications may post on "
                    "official Contoso social media accounts or otherwise communicate as an authorized "
                    "company spokesperson. Employees who are unsure whether a planned post counts as "
                    "speaking on behalf of the company should check with Corporate Communications first.",
                ],
            },
            {
                "heading": "Media and Analyst Inquiries",
                "paragraphs": [
                    "Any inquiry from a journalist, industry analyst, or investor must be redirected to "
                    "Corporate Communications at press@contoso-corp.example rather than answered "
                    "directly, even if the question seems simple or unrelated to sensitive topics.",
                ],
            },
        ],
    },
    {
        "key": "health-safety-policy",
        "title": "Contoso Corp Workplace Health and Safety Policy",
        "doc_type": "pdf",
        "filename": "workplace-health-and-safety-policy.pdf",
        "sections": [
            {
                "heading": "Safety Responsibilities",
                "paragraphs": [
                    "Every employee is responsible for following posted safety procedures and for "
                    "reporting unsafe conditions. Facilities management is responsible for maintaining "
                    "fire suppression systems, emergency lighting, and posted evacuation routes.",
                ],
            },
            {
                "heading": "Incident Reporting",
                "paragraphs": [
                    "Any workplace injury, no matter how minor, must be reported to your manager and to "
                    "Facilities within 24 hours using the incident report form. This applies to injuries "
                    "at any Contoso location, including client sites.",
                ],
            },
            {
                "heading": "Emergency Procedures",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Fire Evacuation",
                        "paragraphs": [
                            "When an alarm sounds, employees must evacuate immediately using the nearest "
                            "marked exit and gather at their floor's designated assembly point. Elevators "
                            "must never be used during a fire alarm.",
                        ],
                    },
                    {
                        "heading": "Medical Emergencies",
                        "paragraphs": [
                            "For a medical emergency, call local emergency services first, then notify "
                            "building security. Trained first-aid responders and AED locations are posted "
                            "on each floor.",
                        ],
                    },
                ],
            },
            {
                "heading": "Workplace Violence Prevention",
                "paragraphs": [
                    "Contoso has zero tolerance for threats or acts of violence in the workplace, "
                    "including threats made by phone, email, or social media. Any concern about "
                    "potential violence should be reported immediately to Security and HR, even if the "
                    "concern feels uncertain.",
                ],
            },
        ],
    },
    {
        "key": "performance-management",
        "title": "Contoso Corp Performance Management and Review Policy",
        "doc_type": "docx",
        "filename": "performance-management-and-review-policy.docx",
        "sections": [
            {
                "heading": "Performance Review Cycle",
                "paragraphs": [
                    "Contoso runs a formal performance review twice per year, in June and December, "
                    "covering the preceding six months of work. Managers and employees are expected to "
                    "have informal check-ins at least monthly outside of the formal cycle.",
                ],
            },
            {
                "heading": "Goal Setting",
                "paragraphs": [
                    "Employees set 3 to 5 goals each review period in partnership with their manager, "
                    "aligned to team priorities. Goals are documented in the performance system and "
                    "revisited at the next check-in if priorities shift.",
                ],
            },
            {
                "heading": "Rating Scale",
                "paragraphs": [
                    "Overall performance is rated on a four-point scale: Exceeds Expectations, Meets "
                    "Expectations, Partially Meets Expectations, and Does Not Meet Expectations. Ratings "
                    "feed into the annual merit increase process described in the Contoso Corp "
                    "Compensation and Pay Practices Policy.",
                ],
            },
            {
                "heading": "Performance Improvement Plans",
                "paragraphs": [
                    "An employee rated Does Not Meet Expectations, or Partially Meets Expectations for "
                    "two consecutive cycles, is placed on a formal Performance Improvement Plan with "
                    "specific, time-bound goals, typically lasting 30 to 60 days.",
                ],
            },
        ],
    },
    {
        "key": "compensation-pay-practices",
        "title": "Contoso Corp Compensation and Pay Practices Policy",
        "doc_type": "pdf",
        "filename": "compensation-and-pay-practices-policy.pdf",
        "sections": [
            {
                "heading": "Pay Grades and Structure",
                "paragraphs": [
                    "Every role at Contoso is assigned to a pay grade based on scope, required skill, "
                    "and market benchmarking, reviewed annually by the Compensation team. Pay grade "
                    "determines the base salary range for a role but not the specific salary within "
                    "that range, which reflects experience and performance.",
                ],
            },
            {
                "heading": "Pay Frequency and Pay Days",
                "paragraphs": [
                    "Employees are paid semi-monthly, on the 15th and last business day of each month, "
                    "by direct deposit. Pay statements are available in the HR system on the morning of "
                    "each pay day.",
                ],
            },
            {
                "heading": "Overtime Eligibility",
                "paragraphs": [
                    "Non-exempt employees are paid 1.5 times their regular hourly rate for hours worked "
                    "beyond 40 in a work week. Overtime must be pre-approved by a manager except in "
                    "genuine emergencies, and unapproved overtime is still paid but may be addressed as a "
                    "performance matter.",
                ],
            },
            {
                "heading": "Annual Merit Increase Process",
                "paragraphs": [
                    "Base salary increases are considered once per year, following the December "
                    "performance review cycle, and are effective with the first full pay period of the "
                    "new calendar year. Increase amounts reflect performance rating, position within the "
                    "pay range, and the year's overall compensation budget.",
                ],
            },
        ],
    },
    {
        "key": "onboarding-guide",
        "title": "Contoso Corp New Hire Onboarding Guide",
        "doc_type": "markdown",
        "filename": "new-hire-onboarding-guide.md",
        "sections": [
            {
                "heading": "Before Your First Day",
                "paragraphs": [
                    "HR sends your offer letter, background check consent, and Day One paperwork through "
                    "the onboarding portal about two weeks before your start date. Complete all forms at "
                    "least 3 business days before you start so your accounts are ready on Day One.",
                ],
            },
            {
                "heading": "Your First Day",
                "paragraphs": [
                    "You will receive your laptop and building access badge from IT, complete benefits "
                    "enrollment elections, and meet with your manager to review your first 30 days. New "
                    "hire orientation covers company overview, the Code of Conduct, and how to navigate "
                    "core HR systems.",
                ],
            },
            {
                "heading": "Your First Week",
                "paragraphs": [
                    "Your manager schedules introductory meetings with your immediate team and key "
                    "cross-functional partners, and assigns an onboarding buddy to answer day-to-day "
                    "questions. You should complete mandatory compliance training, including information "
                    "security awareness, within your first week.",
                ],
            },
            {
                "heading": "30-60-90 Day Milestones",
                "paragraphs": [
                    "At 30 days, you and your manager review early progress against onboarding goals. At "
                    "60 days, you set your first full set of performance goals. At 90 days, your "
                    "introductory period concludes with a formal check-in.",
                ],
            },
        ],
    },
    {
        "key": "offboarding-separation",
        "title": "Contoso Corp Offboarding and Separation Policy",
        "doc_type": "markdown",
        "filename": "offboarding-and-separation-policy.md",
        "sections": [
            {
                "heading": "Types of Separation",
                "paragraphs": [
                    "This policy covers voluntary resignation, involuntary termination, and layoff "
                    "situations for regular employees. Each type follows the same core offboarding "
                    "checklist, with notice-period expectations differing by type as described below.",
                ],
            },
            {
                "heading": "Resignation Notice Period",
                "paragraphs": [
                    "Employees who resign voluntarily are asked to provide at least 2 weeks of written "
                    "notice to their manager, and 4 weeks for people-manager and director-level roles, to "
                    "allow for knowledge transfer. Contoso may, at its discretion, make the separation "
                    "effective before the end of a notice period while still paying through that date.",
                ],
            },
            {
                "heading": "Final Pay and Benefits Continuation",
                "paragraphs": [
                    "Final pay, including any accrued unused PTO, is issued on the next regularly "
                    "scheduled pay date following separation, or sooner where state law requires. Medical "
                    "coverage continues through the end of the month of separation, after which COBRA "
                    "continuation options are offered.",
                ],
            },
            {
                "heading": "Return of Company Property",
                "paragraphs": [
                    "All Contoso equipment, including laptops, badges, and any physical documents, must "
                    "be returned on or before the last working day. IT disables system access at end of "
                    "day on the separation date unless otherwise coordinated with the employee's manager.",
                ],
            },
            {
                "heading": "Exit Interview",
                "paragraphs": [
                    "HR offers a voluntary exit interview to every departing employee to gather feedback "
                    "on their experience. Feedback is aggregated for trends and is not shared with the "
                    "employee's manager in a way that identifies the individual.",
                ],
            },
        ],
    },
    {
        "key": "accommodation-anti-discrimination",
        "title": "Contoso Corp Workplace Accommodation and Anti-Discrimination Policy",
        "doc_type": "docx",
        "filename": "workplace-accommodation-and-anti-discrimination-policy.docx",
        "sections": [
            {
                "heading": "Anti-Discrimination Commitment",
                "paragraphs": [
                    "Contoso prohibits discrimination in any employment decision based on disability, "
                    "religion, or any other characteristic protected by law, and is committed to "
                    "providing reasonable accommodations that enable qualified employees to perform the "
                    "essential functions of their role.",
                ],
            },
            {
                "heading": "Requesting a Disability Accommodation",
                "paragraphs": [],
                "subsections": [
                    {
                        "heading": "Interactive Process",
                        "paragraphs": [
                            "An employee requests an accommodation by contacting HR or their manager; no "
                            "specific words or forms are required to start the process. HR then engages "
                            "in an interactive dialogue with the employee, and where appropriate their "
                            "medical provider, to identify an effective accommodation.",
                        ],
                    },
                    {
                        "heading": "Approval and Implementation",
                        "paragraphs": [
                            "HR aims to respond to an accommodation request within 10 business days. "
                            "Approved accommodations are documented and reviewed periodically to confirm "
                            "they remain effective as job duties or circumstances change.",
                        ],
                    },
                ],
            },
            {
                "heading": "Religious Accommodation",
                "paragraphs": [
                    "Contoso accommodates sincerely held religious beliefs and practices, including "
                    "scheduling adjustments for religious observance and dress or grooming practices, "
                    "unless doing so would create significant operational difficulty.",
                ],
            },
        ],
    },
    {
        "key": "learning-development",
        "title": "Contoso Corp Learning and Professional Development Policy",
        "doc_type": "markdown",
        "filename": "learning-and-professional-development-policy.md",
        "sections": [
            {
                "heading": "Internal Training Catalog",
                "paragraphs": [
                    "Contoso maintains an internal catalog of self-paced and instructor-led courses "
                    "covering technical skills, leadership development, and compliance topics, available "
                    "to all employees at no cost through the Learning Portal.",
                ],
            },
            {
                "heading": "Professional Certification Support",
                "paragraphs": [
                    "Contoso covers the exam fee for one job-relevant professional certification per "
                    "employee per year, subject to manager approval that the certification supports the "
                    "employee's current role or an identified career path.",
                ],
            },
            {
                "heading": "Conference and External Training Attendance",
                "paragraphs": [
                    "Employees may request approval to attend an external conference or training course "
                    "relevant to their role. Approved travel and registration costs follow the "
                    "reimbursement rules in the Contoso Corp Business Travel and Expense Reimbursement "
                    "Policy.",
                ],
            },
            {
                "heading": "Mentorship Program",
                "paragraphs": [
                    "Employees can opt into the Contoso Mentorship Program, which pairs mentees with "
                    "mentors outside their direct reporting line for a 6-month structured cycle focused "
                    "on career development goals.",
                ],
            },
        ],
    },
]


def section_outline(doc):
    outline = []
    for section in doc["sections"]:
        outline.append(
            {
                "heading": section["heading"],
                "subsections": [sub["heading"] for sub in section.get("subsections", [])],
            }
        )
    return outline


def render_markdown(doc):
    lines = [f"# {doc['title']}", ""]
    for section in doc["sections"]:
        lines.append(f"## {section['heading']}")
        lines.append("")
        for para in section.get("paragraphs", []):
            lines.append(para)
            lines.append("")
        for sub in section.get("subsections", []):
            lines.append(f"### {sub['heading']}")
            lines.append("")
            for para in sub.get("paragraphs", []):
                lines.append(para)
                lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def render_docx(doc, path):
    document = DocxDocument()
    document.add_heading(doc["title"], level=0)
    for section in doc["sections"]:
        document.add_heading(section["heading"], level=1)
        for para in section.get("paragraphs", []):
            document.add_paragraph(para)
        for sub in section.get("subsections", []):
            document.add_heading(sub["heading"], level=2)
            for para in sub.get("paragraphs", []):
                document.add_paragraph(para)
    document.save(path)


def render_pdf(doc, path):
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("DocTitle", parent=styles["Title"], fontSize=20, leading=24, spaceAfter=18)
    h1_style = ParagraphStyle(
        "DocH1", parent=styles["Heading1"], fontSize=15, leading=18, spaceBefore=16, spaceAfter=8
    )
    h2_style = ParagraphStyle(
        "DocH2", parent=styles["Heading2"], fontSize=12.5, leading=15, spaceBefore=12, spaceAfter=6
    )
    body_style = ParagraphStyle(
        "DocBody", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=8
    )

    flow = [Paragraph(doc["title"], title_style), Spacer(1, 0.1 * inch)]
    for section in doc["sections"]:
        flow.append(Paragraph(section["heading"], h1_style))
        for para in section.get("paragraphs", []):
            flow.append(Paragraph(para, body_style))
        for sub in section.get("subsections", []):
            flow.append(Paragraph(sub["heading"], h2_style))
            for para in sub.get("paragraphs", []):
                flow.append(Paragraph(para, body_style))

    pdf = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        title=doc["title"],
    )
    pdf.build(flow)


def build_manifest():
    return {
        "corpus": "Contoso Corp Policy Corpus",
        "document_count": len(DOCUMENTS),
        "documents": [
            {
                "filename": doc["filename"],
                "title": doc["title"],
                "type": doc["doc_type"],
                "section_outline": section_outline(doc),
            }
            for doc in DOCUMENTS
        ],
    }


def main():
    for doc in DOCUMENTS:
        path = OUT_DIR / doc["filename"]
        if doc["doc_type"] == "markdown":
            path.write_text(render_markdown(doc), encoding="utf-8")
        elif doc["doc_type"] == "docx":
            render_docx(doc, path)
        elif doc["doc_type"] == "pdf":
            render_pdf(doc, path)
        else:
            raise ValueError(f"Unknown doc_type: {doc['doc_type']}")
        print(f"wrote {path.name}")

    manifest_path = OUT_DIR / "manifest.json"
    manifest_path.write_text(json.dumps(build_manifest(), indent=2) + "\n", encoding="utf-8")
    print(f"wrote {manifest_path.name}")


if __name__ == "__main__":
    main()
